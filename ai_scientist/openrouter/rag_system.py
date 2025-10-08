"""
RAG System for AI-Scientist-v2

Advanced document ingestion, processing, and retrieval system with vector storage.
Supports theory evolution, document chunking, semantic search, and knowledge management.
Enhanced with comprehensive OpenRouter integration and production-grade features.
"""

import os
import hashlib
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple, Union
from dataclasses import dataclass
from datetime import datetime
import asyncio
import aiofiles

# Document processing imports
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    try:
        import PyPDF2
        PYMUPDF_AVAILABLE = False
    except ImportError:
        PyPDF2 = None
        PYMUPDF_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

from bs4 import BeautifulSoup
import requests
from urllib.parse import urlparse

# Local imports
from .config import RAGConfig

# Vector storage and embeddings
try:
    import chromadb
    from chromadb.config import Settings
    CHROMA_AVAILABLE = True
except ImportError:
    CHROMA_AVAILABLE = False

try:
    from sentence_transformers import SentenceTransformer
    SENTENCE_TRANSFORMERS_AVAILABLE = True
except ImportError:
    SENTENCE_TRANSFORMERS_AVAILABLE = False

try:
    import openai
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False

import numpy as np
from sklearn.metrics.pairwise import cosine_similarity

from .utils import chunk_text, count_tokens
import json

logger = logging.getLogger(__name__)

@dataclass
class DocumentMetadata:
    """Metadata for ingested documents"""
    source_file: str
    document_type: str
    file_size: int
    file_hash: str
    ingestion_date: datetime
    chunk_count: int
    total_tokens: int
    embedding_model: str
    tags: List[str]
    custom_fields: Dict[str, Any]

@dataclass
class Document:
    """Document representation"""
    id: str
    title: str
    content: str
    source: str  # file path or URL
    doc_type: str  # pdf, txt, md, docx, url
    metadata: DocumentMetadata
    created_at: datetime
    embedding: Optional[np.ndarray] = None

@dataclass
class DocumentChunk:
    """Document chunk representation"""
    id: str
    doc_id: str
    content: str
    chunk_index: int
    start_char: int
    end_char: int
    embedding: Optional[np.ndarray] = None
    metadata: Dict[str, Any] = None

class DocumentProcessor:
    """Enhanced document processing utilities with support for multiple formats"""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.txt': self._process_text,
            '.md': self._process_text,
            '.docx': self._process_docx,
            '.json': self._process_json,
            '.csv': self._process_csv,
            '.py': self._process_text,
            '.ipynb': self._process_jupyter,
            '.html': self._process_html,
            '.xml': self._process_xml
        }
    
    def is_supported(self, file_path: Path) -> bool:
        """Check if file format is supported"""
        return file_path.suffix.lower() in self.supported_formats
    
    async def process_file(self, file_path: Path) -> Tuple[str, DocumentMetadata]:
        """Process a file and extract text content with metadata"""
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if not self.is_supported(file_path):
            raise ValueError(f"Unsupported file format: {file_path.suffix}")
        
        # Calculate file hash for deduplication
        file_hash = self._calculate_file_hash(file_path)
        
        # Process file based on format
        processor = self.supported_formats[file_path.suffix.lower()]
        text_content = await processor(file_path)
        
        # Create metadata
        file_stats = file_path.stat()
        metadata = DocumentMetadata(
            source_file=str(file_path),
            document_type=file_path.suffix.lower(),
            file_size=file_stats.st_size,
            file_hash=file_hash,
            ingestion_date=datetime.now(),
            chunk_count=0,  # Will be updated after chunking
            total_tokens=count_tokens(text_content),
            embedding_model="",  # Will be updated by RAG system
            tags=[],
            custom_fields={}
        )
        
        logger.info(f"Processed {file_path.name}: {len(text_content)} chars, {metadata.total_tokens} tokens")
        return text_content, metadata
    
    def _calculate_file_hash(self, file_path: Path) -> str:
        """Calculate SHA-256 hash of file"""
        hash_sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()
    
    async def _process_pdf(self, file_path: Path) -> str:
        """Extract text from PDF files using PyMuPDF (preferred) or PyPDF2"""
        try:
            if PYMUPDF_AVAILABLE:
                doc = fitz.open(file_path)
                text_content = []
                
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    text = page.get_text()
                    if text.strip():
                        text_content.append(f"--- Page {page_num + 1} ---\n{text}")
                
                doc.close()
                return "\n\n".join(text_content)
            elif PyPDF2:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                    return text.strip()
            else:
                raise ImportError("No PDF processing library available")
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            raise
    
    async def _process_text(self, file_path: Path) -> str:
        """Process plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
    
    async def _process_docx(self, file_path: Path) -> str:
        """Extract text from DOCX files"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX support")
        
        try:
            doc = docx.Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs)
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            raise
    
    async def _process_json(self, file_path: Path) -> str:
        """Process JSON files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            return json.dumps(data, indent=2, ensure_ascii=False)
        except Exception as e:
            logger.error(f"Error processing JSON {file_path}: {e}")
            raise
    
    async def _process_csv(self, file_path: Path) -> str:
        """Process CSV files"""
        if not PANDAS_AVAILABLE:
            # Fallback to basic CSV reading
            try:
                import csv
                with open(file_path, 'r', encoding='utf-8') as f:
                    reader = csv.reader(f)
                    rows = list(reader)
                    return "\n".join([",".join(row) for row in rows])
            except Exception as e:
                logger.error(f"Error processing CSV {file_path}: {e}")
                raise
        
        try:
            df = pd.read_csv(file_path)
            return f"CSV Data Summary:\nColumns: {', '.join(df.columns)}\nRows: {len(df)}\n\nData:\n{df.to_string()}"
        except Exception as e:
            logger.error(f"Error processing CSV {file_path}: {e}")
            raise
    
    async def _process_jupyter(self, file_path: Path) -> str:
        """Extract text and code from Jupyter notebooks"""
        try:
            import json
            with open(file_path, 'r', encoding='utf-8') as f:
                notebook = json.load(f)
            
            content_parts = []
            for i, cell in enumerate(notebook.get('cells', [])):
                cell_type = cell.get('cell_type', 'unknown')
                source = ''.join(cell.get('source', []))
                
                if source.strip():
                    content_parts.append(f"--- Cell {i+1} ({cell_type}) ---\n{source}")
                
                # Include outputs for code cells
                if cell_type == 'code' and 'outputs' in cell:
                    for j, output in enumerate(cell['outputs']):
                        if 'text' in output:
                            output_text = ''.join(output['text'])
                            content_parts.append(f"--- Cell {i+1} Output {j+1} ---\n{output_text}")
            
            return "\n\n".join(content_parts)
        except Exception as e:
            logger.error(f"Error processing Jupyter notebook {file_path}: {e}")
            raise
    
    async def _process_html(self, file_path: Path) -> str:
        """Process HTML files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            soup = BeautifulSoup(content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text content
            return soup.get_text(separator='\n', strip=True)
        except Exception as e:
            logger.error(f"Error processing HTML {file_path}: {e}")
            raise
    
    async def _process_xml(self, file_path: Path) -> str:
        """Process XML files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            soup = BeautifulSoup(content, 'xml')
            return soup.get_text(separator='\n', strip=True)
        except Exception as e:
            logger.error(f"Error processing XML {file_path}: {e}")
            raise
    
    @staticmethod
    def extract_text_from_txt(file_path: Path) -> str:
        """Extract text from TXT/MD file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {e}")
            return ""
    
    @staticmethod
    def extract_text_from_url(url: str) -> Tuple[str, str]:
        """Extract text from URL"""
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            response = requests.get(url, headers=headers, timeout=30)
            response.raise_for_status()
            
            if 'application/pdf' in response.headers.get('content-type', ''):
                # Handle PDF URLs
                with open('/tmp/temp_pdf.pdf', 'wb') as f:
                    f.write(response.content)
                text = DocumentProcessor.extract_text_from_pdf(Path('/tmp/temp_pdf.pdf'))
                os.remove('/tmp/temp_pdf.pdf')
                title = urlparse(url).path.split('/')[-1]
            else:
                # Handle HTML URLs
                soup = BeautifulSoup(response.content, 'html.parser')
                
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                
                # Get title
                title_tag = soup.find('title')
                title = title_tag.get_text().strip() if title_tag else urlparse(url).netloc
                
                # Get main content
                text = soup.get_text(separator='\n', strip=True)
            
            return text, title
            
        except Exception as e:
            logger.error(f"Error extracting text from URL {url}: {e}")
            return "", ""

    @staticmethod
    def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """Split text into overlapping chunks"""
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to end at a sentence boundary
            if end < len(text):
                # Look for sentence endings within the last 100 characters
                sentence_ends = ['. ', '! ', '? ', '\n\n']
                for i in range(min(100, end - start)):
                    pos = end - i
                    if any(text[pos:pos+2] == ending for ending in sentence_ends):
                        end = pos + 1
                        break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = max(start + 1, end - overlap)
            if start >= len(text):
                break
        
        return chunks

class EmbeddingGenerator:
    """Generate embeddings for documents and queries"""
    
    def __init__(self, model_name: str = "text-embedding-3-large"):
        self.model_name = model_name
        self.openai_client = None
        self.sentence_transformer = None
        
        if OPENAI_AVAILABLE and model_name.startswith("text-embedding"):
            self.openai_client = openai.OpenAI()
        else:
            # Fallback to sentence transformers
            try:
                self.sentence_transformer = SentenceTransformer('all-MiniLM-L6-v2')
                logger.info("Using SentenceTransformer for embeddings")
            except Exception as e:
                logger.error(f"Error loading sentence transformer: {e}")
    
    def generate_embedding(self, text: str) -> np.ndarray:
        """Generate embedding for text"""
        try:
            if self.openai_client and self.model_name.startswith("text-embedding"):
                response = self.openai_client.embeddings.create(
                    model=self.model_name,
                    input=text
                )
                return np.array(response.data[0].embedding)
            elif self.sentence_transformer:
                return self.sentence_transformer.encode(text)
            else:
                raise ValueError("No embedding model available")
        except Exception as e:
            logger.error(f"Error generating embedding: {e}")
            # Return zero vector as fallback
            return np.zeros(384)  # Default dimension for MiniLM
    
    def generate_batch_embeddings(self, texts: List[str]) -> List[np.ndarray]:
        """Generate embeddings for multiple texts"""
        try:
            if self.openai_client and self.model_name.startswith("text-embedding"):
                response = self.openai_client.embeddings.create(
                    model=self.model_name,
                    input=texts
                )
                return [np.array(item.embedding) for item in response.data]
            elif self.sentence_transformer:
                embeddings = self.sentence_transformer.encode(texts)
                return [embedding for embedding in embeddings]
            else:
                raise ValueError("No embedding model available")
        except Exception as e:
            logger.error(f"Error generating batch embeddings: {e}")
            return [np.zeros(384) for _ in texts]

class ChromaVectorStore:
    """Chroma-based vector storage"""
    
    def __init__(self, config: RAGConfig, persist_directory: str = "./chroma_db"):
        self.config = config
        self.persist_directory = Path(persist_directory)
        self.persist_directory.mkdir(parents=True, exist_ok=True)
        
        if not CHROMA_AVAILABLE:
            raise ImportError("ChromaDB is required but not installed")
        
        self.client = chromadb.PersistentClient(
            path=str(self.persist_directory),
            settings=Settings(allow_reset=True)
        )
        
        # Create collection
        self.collection_name = "ai_scientist_docs"
        try:
            self.collection = self.client.get_collection(name=self.collection_name)
        except:
            self.collection = self.client.create_collection(
                name=self.collection_name,
                metadata={"hnsw:space": "cosine"}
            )
    
    def add_chunks(self, chunks: List[DocumentChunk]) -> None:
        """Add document chunks to the vector store"""
        try:
            # Prepare data for ChromaDB
            ids = [chunk.id for chunk in chunks]
            documents = [chunk.content for chunk in chunks]
            metadatas = []
            embeddings = []
            
            for chunk in chunks:
                metadata = {
                    "doc_id": chunk.doc_id,
                    "chunk_index": chunk.chunk_index,
                    "start_char": chunk.start_char,
                    "end_char": chunk.end_char,
                    **(chunk.metadata or {})
                }
                metadatas.append(metadata)
                
                if chunk.embedding is not None:
                    embeddings.append(chunk.embedding.tolist())
            
            if embeddings:
                self.collection.add(
                    ids=ids,
                    documents=documents,
                    metadatas=metadatas,
                    embeddings=embeddings
                )
            else:
                # Let ChromaDB generate embeddings
                self.collection.add(
                    ids=ids,
                    documents=documents,
                    metadatas=metadatas
                )
            
            logger.info(f"Added {len(chunks)} chunks to vector store")
            
        except Exception as e:
            logger.error(f"Error adding chunks to vector store: {e}")
    
    def search(self, query: str, n_results: int = 10) -> List[Tuple[str, float, Dict[str, Any]]]:
        """Search for similar documents"""
        try:
            results = self.collection.query(
                query_texts=[query],
                n_results=n_results,
                include=['documents', 'metadatas', 'distances']
            )
            
            search_results = []
            for i in range(len(results['ids'][0])):
                content = results['documents'][0][i]
                distance = results['distances'][0][i]
                metadata = results['metadatas'][0][i]
                similarity = 1 - distance  # Convert distance to similarity
                
                search_results.append((content, similarity, metadata))
            
            return search_results
            
        except Exception as e:
            logger.error(f"Error searching vector store: {e}")
            return []
    
    def delete_document(self, doc_id: str) -> None:
        """Delete all chunks for a document"""
        try:
            # Get all chunks for the document
            results = self.collection.get(
                where={"doc_id": doc_id},
                include=['ids']
            )
            
            if results['ids']:
                self.collection.delete(ids=results['ids'])
                logger.info(f"Deleted document {doc_id} from vector store")
        except Exception as e:
            logger.error(f"Error deleting document {doc_id}: {e}")

class RAGSystem:
    """Complete RAG system for document ingestion and retrieval"""
    
    def __init__(self, config: RAGConfig, storage_dir: str = "./rag_storage"):
        self.config = config
        self.storage_dir = Path(storage_dir)
        self.storage_dir.mkdir(parents=True, exist_ok=True)
        
        # Initialize components
        self.document_processor = DocumentProcessor()
        self.embedding_generator = EmbeddingGenerator(config.embedding_model)
        
        # Initialize vector store
        if config.vector_store == "chroma":
            self.vector_store = ChromaVectorStore(config, str(self.storage_dir / "chroma"))
        else:
            raise ValueError(f"Vector store {config.vector_store} not supported")
        
        # Document registry
        self.documents_file = self.storage_dir / "documents.json"
        self.documents: Dict[str, Document] = self._load_documents()
    
    def _load_documents(self) -> Dict[str, Document]:
        """Load document registry"""
        try:
            if self.documents_file.exists():
                import json
                with open(self.documents_file, 'r') as f:
                    docs_data = json.load(f)
                    
                documents = {}
                for doc_id, doc_data in docs_data.items():
                    documents[doc_id] = Document(
                        id=doc_data['id'],
                        title=doc_data['title'],
                        content=doc_data['content'],
                        source=doc_data['source'],
                        doc_type=doc_data['doc_type'],
                        metadata=doc_data['metadata'],
                        created_at=datetime.fromisoformat(doc_data['created_at'])
                    )
                
                return documents
        except Exception as e:
            logger.error(f"Error loading documents: {e}")
        
        return {}
    
    def _save_documents(self) -> None:
        """Save document registry"""
        try:
            import json
            docs_data = {}
            for doc_id, doc in self.documents.items():
                docs_data[doc_id] = {
                    'id': doc.id,
                    'title': doc.title,
                    'content': doc.content,
                    'source': doc.source,
                    'doc_type': doc.doc_type,
                    'metadata': doc.metadata,
                    'created_at': doc.created_at.isoformat()
                }
            
            with open(self.documents_file, 'w') as f:
                json.dump(docs_data, f, indent=2)
                
        except Exception as e:
            logger.error(f"Error saving documents: {e}")
    
    def _generate_document_id(self, source: str, content: str) -> str:
        """Generate unique document ID"""
        content_hash = hashlib.md5(content.encode()).hexdigest()
        source_hash = hashlib.md5(source.encode()).hexdigest()
        return f"{source_hash[:8]}_{content_hash[:8]}"
    
    async def ingest_file(self, file_path: Path, metadata: Optional[Dict[str, Any]] = None) -> Optional[str]:
        """Ingest a file into the RAG system"""
        try:
            if not file_path.exists():
                logger.error(f"File not found: {file_path}")
                return None
            
            # Process file using document processor
            content, doc_metadata = await self.document_processor.process_file(file_path)
            
            if not content:
                logger.error(f"No content extracted from {file_path}")
                return None
            
            # Create document
            doc_id = self._generate_document_id(str(file_path), content)
            title = metadata.get('title', file_path.stem) if metadata else file_path.stem
            
            document = Document(
                id=doc_id,
                title=title,
                content=content,
                source=str(file_path),
                doc_type=doc_metadata.document_type,
                metadata=doc_metadata,
                created_at=datetime.now()
            )
            
            return await self._ingest_document(document)
            
        except Exception as e:
            logger.error(f"Error ingesting file {file_path}: {e}")
            return None
    
    def ingest_url(self, url: str, metadata: Optional[Dict[str, Any]] = None) -> Optional[str]:
        """Ingest content from a URL"""
        try:
            content, title = self.document_processor.extract_text_from_url(url)
            
            if not content:
                logger.error(f"No content extracted from URL: {url}")
                return None
            
            # Create document
            doc_id = self._generate_document_id(url, content)
            title = metadata.get('title', title) if metadata else title
            
            document = Document(
                id=doc_id,
                title=title,
                content=content,
                source=url,
                doc_type='url',
                metadata=metadata or {},
                created_at=datetime.now()
            )
            
            return self._ingest_document(document)
            
        except Exception as e:
            logger.error(f"Error ingesting URL {url}: {e}")
            return None
    
    async def _ingest_document(self, document: Document) -> str:
        """Internal method to ingest a document"""
        try:
            # Check if document already exists
            if document.id in self.documents:
                logger.info(f"Document {document.id} already exists, skipping")
                return document.id
            
            # Split into chunks
            chunks_text = self.document_processor.chunk_text(
                document.content,
                self.config.chunk_size,
                self.config.chunk_overlap
            )
            
            # Generate embeddings for chunks
            embeddings = self.embedding_generator.generate_batch_embeddings(chunks_text)
            
            # Create document chunks
            chunks = []
            for i, (chunk_text, embedding) in enumerate(zip(chunks_text, embeddings)):
                chunk_id = f"{document.id}_chunk_{i}"
                
                # Calculate character positions
                start_char = document.content.find(chunk_text[:50])  # Approximate
                end_char = start_char + len(chunk_text)
                
                chunk = DocumentChunk(
                    id=chunk_id,
                    doc_id=document.id,
                    content=chunk_text,
                    chunk_index=i,
                    start_char=max(0, start_char),
                    end_char=end_char,
                    embedding=embedding,
                    metadata={
                        "title": document.title,
                        "source": document.source,
                        "doc_type": document.doc_type,
                        **document.metadata
                    }
                )
                chunks.append(chunk)
            
            # Add to vector store
            self.vector_store.add_chunks(chunks)
            
            # Save document
            self.documents[document.id] = document
            self._save_documents()
            
            logger.info(f"Successfully ingested document: {document.title} ({len(chunks)} chunks)")
            return document.id
            
        except Exception as e:
            logger.error(f"Error ingesting document {document.title}: {e}")
            return None
    
    def search(self, query: str, max_results: Optional[int] = None) -> List[Tuple[str, float, Dict[str, Any]]]:
        """Search for relevant documents"""
        max_results = max_results or self.config.max_retrieval_docs
        
        # Search vector store
        results = self.vector_store.search(query, max_results)
        
        # Filter by similarity threshold
        filtered_results = [
            (content, score, metadata)
            for content, score, metadata in results
            if score >= self.config.similarity_threshold
        ]
        
        logger.info(f"Found {len(filtered_results)} relevant chunks for query")
        return filtered_results
    
    def get_context_for_query(self, query: str, max_context_length: int = 4000) -> str:
        """Get formatted context for a query"""
        results = self.search(query)
        
        if not results:
            return ""
        
        # Build context from top results
        context_parts = []
        current_length = 0
        
        for content, score, metadata in results:
            # Format context entry
            source = metadata.get('source', 'Unknown')
            title = metadata.get('title', 'Untitled')
            
            context_entry = f"Source: {title} ({source})\nContent: {content}\n"
            
            if current_length + len(context_entry) > max_context_length:
                break
            
            context_parts.append(context_entry)
            current_length += len(context_entry)
        
        context = "\n---\n".join(context_parts)
        
        if context:
            context = f"RELEVANT CONTEXT:\n{context}\n---\n"
        
        return context
    
    def delete_document(self, doc_id: str) -> bool:
        """Delete a document and its chunks"""
        try:
            if doc_id not in self.documents:
                logger.warning(f"Document {doc_id} not found")
                return False
            
            # Delete from vector store
            self.vector_store.delete_document(doc_id)
            
            # Delete from registry
            del self.documents[doc_id]
            self._save_documents()
            
            logger.info(f"Deleted document {doc_id}")
            return True
            
        except Exception as e:
            logger.error(f"Error deleting document {doc_id}: {e}")
            return False
    
    def list_documents(self) -> List[Dict[str, Any]]:
        """List all ingested documents"""
        return [
            {
                'id': doc.id,
                'title': doc.title,
                'source': doc.source,
                'doc_type': doc.doc_type,
                'created_at': doc.created_at.isoformat(),
                'content_length': len(doc.content),
                'metadata': doc.metadata
            }
            for doc in self.documents.values()
        ]
    
    def get_statistics(self) -> Dict[str, Any]:
        """Get RAG system statistics"""
        total_docs = len(self.documents)
        total_content_length = sum(len(doc.content) for doc in self.documents.values())
        
        doc_types = {}
        for doc in self.documents.values():
            doc_types[doc.doc_type] = doc_types.get(doc.doc_type, 0) + 1
        
        return {
            'total_documents': total_docs,
            'total_content_length': total_content_length,
            'document_types': doc_types,
            'vector_store': self.config.vector_store,
            'embedding_model': self.config.embedding_model,
            'chunk_size': self.config.chunk_size,
            'chunk_overlap': self.config.chunk_overlap
        }